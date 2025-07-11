from __future__ import annotations

import os
import csv
import json
from pathlib import Path
from typing import List

MAX_PROMPT_CHARS = 25000

from django.conf import settings
from django.core.files.base import File

import openai
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from PyPDF2 import PdfReader



def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    """Insert a new paragraph after the given one and return it.

    If a style is provided but not found in the document, the paragraph will be
    added using the document's default style instead of raising ``KeyError``.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except KeyError:
            # Fall back to default style if the requested style is not present
            pass
    return new_para


from .models import AIResult, ContextFile, ExamFile
from config.utils import load_prompts




def _read_file(path: str) -> str:
    """Return text content from a file path."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if ext == ".csv":
        with open(path, newline="", encoding="utf-8", errors="ignore") as fh:
            reader = csv.reader(fh)
            return "\n".join(",".join(row) for row in reader)
    if ext == ".pdf":
        try:
            with open(path, "rb") as fh:
                reader = PdfReader(fh)
                return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            return ""
    # Fallback to plain text
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        return fh.read()


def assemble_prompt(session_id: str) -> str:
    """Create a base prompt from uploaded exam and context files."""
    exam = ExamFile.objects.filter(session_id=session_id).first()
    if not exam:
        raise ValueError("No exam file uploaded")
    context_files = ContextFile.objects.filter(session_id=session_id)

    exam_text = _read_file(exam.file.path)
    context_text = "\n\n".join(_read_file(c.file.path) for c in context_files)

    prompt = f"Exam:\n{exam_text}\n\nAdditional context:\n{context_text}"
    if len(prompt) > MAX_PROMPT_CHARS:
        raise ValueError("Zu viele oder zu große Kontextdateien f\u00fcr die KI")
    return prompt


def _insert_answers_ai(
    doc: Document,
    answer: str,
    level: str,
    color: RGBColor,
    client: openai.OpenAI,
    model: str,
) -> bool:
    """Use the AI to find insertion spots for an answer within a document.

    The AI receives the numbered exam paragraphs and the student's complete
    answer. It returns JSON objects with the paragraph index after which each
    answer segment should be inserted. The function inserts these segments and
    returns ``True`` if at least one insertion was made."""

    numbered = "\n".join(f"{i}: {p.text}" for i, p in enumerate(doc.paragraphs))
    system_msg = (
        "You receive an exam with numbered paragraphs and a student's answer."
        " Decide after which paragraphs the answer segments belong and return a"
        " JSON list of objects with keys 'after' (paragraph number) and 'text'"
        " (exact answer text). Do not change wording or language."
    )
    user_msg = (
        f"Exam paragraphs:\n{numbered}\n\nStudent answer:\n{answer}\n\n"
        "Return only the JSON list."
    )

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
            temperature=0,
        )
        instructions = json.loads(resp.choices[0].message.content)
    except Exception:
        return False

    if not isinstance(instructions, list):
        return False

    made_change = False
    for ins in sorted(instructions, key=lambda x: x.get("after", -1), reverse=True):
        try:
            idx = int(ins.get("after"))
            text = str(ins.get("text", "")).strip()
        except Exception:
            continue
        if not text or idx < 0 or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        head = _insert_paragraph_after(para, f"{level.title()} Antwort:", style="Heading2")
        ans_p = _insert_paragraph_after(head, "")
        run = ans_p.add_run(text)
        run.font.color.rgb = color
        made_change = True

    return made_change


def generate_ai_results(session_id: str, *, api_key: str | None = None) -> List[AIResult]:
    """Generate AI answers for all performance levels."""
    base_prompt = assemble_prompt(session_id)
    client = openai.OpenAI(api_key=api_key or os.environ.get("OPENAI_API_KEY"))

    storage_root = getattr(settings, "MEDIA_ROOT", "")
    session_dir = Path(storage_root) / "ai_results" / session_id
    session_dir.mkdir(parents=True, exist_ok=True)

    # Remove old results for this session
    for old in AIResult.objects.filter(session_id=session_id):
        old.file.delete(save=False)
        old.delete()

    # Get the uploaded exam document to use as template
    exam = ExamFile.objects.filter(session_id=session_id).first()
    if not exam:
        raise ValueError("No exam file uploaded")

    prompts = load_prompts()
    results: List[AIResult] = []
    level_map = {
        "low": prompts["level_low"],
        "medium": prompts["level_medium"],
        "high": prompts["level_high"],
    }
    color_map = {
        "low": RGBColor(0x80, 0x00, 0x00),
        "medium": RGBColor(0x00, 0x64, 0x00),
        "high": RGBColor(0x00, 0x00, 0x80),
    }

    for level, instruction in level_map.items():
        messages = [
            {"role": "system", "content": prompts["system"]},
            {
                "role": "user",
                "content": f"{base_prompt}\n\n{prompts['base']}\n{instruction}",
            },
        ]
        model = getattr(settings, "OPENAI_MODEL", "gpt-4-1106-preview")
        response = client.chat.completions.create(model=model, messages=messages)
        answer = response.choices[0].message.content.strip()

        # Load exam document fresh for each level
        doc = Document(exam.file.path)
        inserted = False
        for para in doc.paragraphs:
            if "[Antwort]" in para.text:
                para.text = para.text.replace("[Antwort]", "").rstrip()
                head = _insert_paragraph_after(
                    para, f"{level.title()} Antwort:", style="Heading2"
                )
                ans_p = _insert_paragraph_after(head, "")
                run = ans_p.add_run(answer)
                run.font.color.rgb = color_map[level]
                inserted = True

        if not inserted:
            inserted = _insert_answers_ai(
                doc,
                answer,
                level,
                color_map[level],
                client,
                model,
            )

        if not inserted:
            # Add heading with graceful fallback if the "Heading 2" style is
            # missing in the template
            try:
                head = doc.add_heading(f"{level.title()} Antwort:", level=2)
            except KeyError:
                head = doc.add_paragraph(f"{level.title()} Antwort:")
                try:
                    head.style = "Heading2"
                except KeyError:
                    pass
            p = doc.add_paragraph()
            run = p.add_run(answer)
            run.font.color.rgb = color_map[level]

        orig_name = Path(exam.file.name).stem
        file_name = f"{orig_name}_{level}.docx"
        file_path = session_dir / file_name
        doc.save(file_path)

        result = AIResult(level=level, session_id=session_id)
        with open(file_path, "rb") as fh:
            result.file.save(f"{session_id}/{file_name}", File(fh), save=True)
        results.append(result)

    return results
